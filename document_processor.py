# -*- coding: utf-8 -*-
"""
文档处理工具模块（轻量版）
支持读取 PDF、Word、TXT 等文档格式
针对服务器内存优化：不使用本地向量模型，使用关键词匹配
"""
import os
import io
import hashlib
import json
from typing import List, Dict, Any, Optional

# 文档处理
try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# 文本分割
try:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    LANGCHAIN_AVAILABLE = True
except ImportError:
    LANGCHAIN_AVAILABLE = False


class DocumentLoader:
    """文档加载器基类"""
    
    @staticmethod
    def load_txt(file_content: bytes, encoding: str = 'utf-8') -> str:
        """加载TXT文本"""
        return file_content.decode(encoding)
    
    @staticmethod
    def load_pdf(file_content: bytes) -> str:
        """加载PDF文档"""
        if not PDF_AVAILABLE:
            raise ImportError("请安装 PyPDF2: pip install PyPDF2")
        
        text = ""
        pdf_file = io.BytesIO(file_content)
        reader = PdfReader(pdf_file)
        
        for page in reader.pages:
            text += page.extract_text() + "\n"
        
        return text
    
    @staticmethod
    def load_docx(file_content: bytes) -> str:
        """加载Word文档"""
        if not DOCX_AVAILABLE:
            raise ImportError("请安装 python-docx: pip install python-docx")
        
        doc_file = io.BytesIO(file_content)
        doc = Document(doc_file)
        
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text
    
    @classmethod
    def load_document(cls, file_content: bytes, file_type: str) -> str:
        """根据文件类型加载文档"""
        file_type = file_type.lower()
        
        if file_type in ['.txt', 'text']:
            return cls.load_txt(file_content)
        elif file_type in ['.pdf', 'pdf']:
            return cls.load_pdf(file_content)
        elif file_type in ['.docx', '.doc', 'word']:
            return cls.load_docx(file_content)
        else:
            raise ValueError(f"不支持的文件类型: {file_type}")


class TextSplitter:
    """文本分割器"""
    
    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 100):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def split_text(self, text: str) -> List[str]:
        """分割文本为块"""
        if LANGCHAIN_AVAILABLE:
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap,
                separators=["\n\n", "\n", "。", "！", "？", " ", ""]
            )
            return splitter.split_text(text)
        else:
            # 简单的备用分割方法
            return self._simple_split(text)
    
    def _simple_split(self, text: str) -> List[str]:
        """简单分割方法"""
        chunks = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            end = min(start + self.chunk_size, text_length)
            chunk = text[start:end]
            chunks.append(chunk)
            start = end - self.chunk_overlap
        
        return chunks


class DocumentProcessor:
    """文档处理器 - 完整流程"""
    
    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 100):
        self.loader = DocumentLoader()
        self.splitter = TextSplitter(chunk_size, chunk_overlap)
    
    def process_document(self, file_content: bytes, file_type: str, 
                        file_name: str = "document") -> Dict[str, Any]:
        """处理文档的完整流程"""
        # 1. 加载文档
        text = self.loader.load_document(file_content, file_type)
        total_chars = len(text)
        
        # 2. 分割文本
        chunks = self.splitter.split_text(text)
        
        # 3. 构建结果
        result_chunks = []
        for i, chunk in enumerate(chunks):
            result_chunks.append({
                "index": i,
                "content": chunk,
                "char_count": len(chunk)
            })
        
        # 4. 生成文档ID
        doc_id = self._generate_doc_id(file_name, text)
        
        return {
            "doc_id": doc_id,
            "file_name": file_name,
            "file_type": file_type,
            "total_chars": total_chars,
            "total_chunks": len(chunks),
            "chunks": result_chunks
        }
    
    def _generate_doc_id(self, file_name: str, content: str) -> str:
        """生成文档唯一ID"""
        unique_str = f"{file_name}_{len(content)}"
        return hashlib.md5(unique_str.encode()).hexdigest()[:16]


class EmbeddingProcessor:
    """
    轻量级向量化处理器
    使用关键词匹配+TF-IDF，无需加载本地向量模型，节省内存
    """
    
    def __init__(self):
        self.embeddings = None
        self.vocab = {}  # 词汇表
        self.use_embedding = False
        print("[OK] 使用轻量级关键词匹配嵌入")
    
    def _build_vocab(self, texts: List[str]):
        """构建词汇表"""
        import re
        self.vocab = {}
        for text in texts:
            # 提取中文词
            words = re.findall(r'[\u4e00-\u9fff]{2,}', text)
            for word in words:
                self.vocab[word] = self.vocab.get(word, 0) + 1
        
        # 取top 1000词汇
        self.vocab = dict(sorted(self.vocab.items(), key=lambda x: x[1], reverse=True)[:1000])
        self.use_embedding = len(self.vocab) > 0
    
    def _text_to_vector(self, text: str) -> List[float]:
        """将文本转为向量（基于词频）"""
        if not self.vocab:
            return [0.0] * len(self.vocab) if self.vocab else [0.0]
        
        import re
        words = re.findall(r'[\u4e00-\u9fff]{2,}', text)
        vector = [0.0] * len(self.vocab)
        
        for i, word in enumerate(self.vocab.keys()):
            vector[i] = words.count(word)
        
        return vector
    
    def embed_texts(self, texts: List[str]) -> List[List[float]]:
        """将文本转为向量"""
        if not self.use_embedding:
            self._build_vocab(texts)
        
        return [self._text_to_vector(text) for text in texts]
    
    def embed_query(self, query: str) -> List[float]:
        """将查询转为向量"""
        return self._text_to_vector(query)


class KeywordVectorStore:
    """
    轻量级向量存储
    使用关键词匹配+倒排索引，不存储大量向量，节省内存
    """
    
    def __init__(self):
        self.documents = []  # 存储文档块
        self.keywords = {}   # 倒排索引: keyword -> [doc_indices]
        self.metadata = []   # 存储元数据
    
    def add_documents(self, chunks: List[str], metadata: Dict[str, Any]):
        """添加文档到存储（不生成向量）"""
        import re
        
        for i, chunk in enumerate(chunks):
            # 提取关键词
            words = re.findall(r'[\u4e00-\u9fff]{2,}', chunk)
            
            # 去重
            unique_words = list(set(words))
            
            self.documents.append(chunk)
            self.metadata.append({**metadata, "chunk_index": i})
            
            # 更新倒排索引
            for word in unique_words:
                if word not in self.keywords:
                    self.keywords[word] = []
                if i not in self.keywords[word]:
                    self.keywords[word].append(i)
    
    def similarity_search(self, query: str, top_k: int = 3) -> List[Dict]:
        """基于关键词的相似度搜索"""
        import re
        
        # 提取查询词
        query_words = re.findall(r'[\u4e00-\u9fff]{2,}', query)
        
        if not query_words or not self.keywords:
            return []
        
        # 计算每个文档的得分
        scores = {}
        for word in query_words:
            if word in self.keywords:
                for doc_idx in self.keywords[word]:
                    scores[doc_idx] = scores.get(doc_idx, 0) + 1
        
        # 排序
        sorted_scores = sorted(scores.items(), key=lambda x: x[1], reverse=True)
        
        # 返回top_k
        results = []
        for idx, score in sorted_scores[:top_k]:
            results.append({
                "content": self.documents[idx],
                "score": score / len(query_words),  # 归一化分数
                "metadata": self.metadata[idx]
            })
        
        return results
    
    def save_to_file(self, filepath: str):
        """保存到文件"""
        data = {
            "documents": self.documents,
            "metadata": self.metadata,
            "keywords": self.keywords
        }
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    
    def load_from_file(self, filepath: str):
        """从文件加载"""
        with open(filepath, 'r', encoding='utf-8') as f:
            data = json.load(f)
            self.documents = data.get("documents", [])
            self.metadata = data.get("metadata", [])
            self.keywords = data.get("keywords", {})


# 兼容旧代码
VectorStore = KeywordVectorStore


# 测试代码
if __name__ == "__main__":
    # 测试文本分割
    test_text = """
    这是一个测试文档。它包含多个段落。
    
    这是第二个段落的内容。我们可以从中提取有用的信息。
    
    这是第三个段落，展示了文档处理的强大功能。
    """
    
    processor = DocumentProcessor(chunk_size=100, chunk_overlap=20)
    
    # 模拟文件处理
    result = processor.process_document(
        test_text.encode('utf-8'),
        '.txt',
        'test.txt'
    )
    
    print("文档处理结果:")
    print(f"总字符数: {result['total_chars']}")
    print(f"总块数: {result['total_chunks']}")
    print("\n文本块:")
    for chunk in result['chunks']:
        print(f"  块{chunk['index']}: {chunk['content'][:50]}...")
